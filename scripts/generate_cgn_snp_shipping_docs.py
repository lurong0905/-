#!/usr/bin/env python
# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import json
import re
import shutil
from collections import OrderedDict, defaultdict
from copy import deepcopy
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side


SUPPLIER = "江苏道众能源科技有限公司"
SUPPLIER_ADDRESS = "江苏省连云港市海州区联东U谷A5-2"
SUPPLIER_PHONE = "0518-81180066"


def clean(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return re.sub(r"\s+", " ", str(value).replace("\r", "\n").strip())


def split_multivalue(value) -> list[str]:
    if value is None:
        return []
    return [clean(part) for part in re.split(r"[|\r\n]+", str(value)) if clean(part)]


def parse_boxes(value: str) -> list[int]:
    text = clean(value)
    match = re.fullmatch(r"(\d+)\s*[-~至]\s*(\d+)", text)
    if match:
        return list(range(int(match.group(1)), int(match.group(2)) + 1))
    return [int(item) for item in re.findall(r"\d+", text)]


def split_counted_parts(value: str, box_count: int) -> list[str]:
    parts = split_multivalue(value)
    if not parts:
        return [""] * box_count
    expanded: list[str] = []
    for part in parts:
        count = 1
        match = re.fullmatch(r"(.+?)\*(\d+)\s*个", part)
        if match:
            part = clean(match.group(1))
            count = int(match.group(2))
        expanded.extend([part] * count)
    if len(expanded) == 1:
        expanded = expanded * box_count
    if len(expanded) < box_count:
        expanded.extend([expanded[-1]] * (box_count - len(expanded)))
    return expanded[:box_count]


def split_weight_parts(value: str, size_value: str, box_count: int) -> list[str]:
    parts = split_multivalue(value)
    if not parts:
        return [""] * box_count
    if len(parts) == 1:
        return parts * box_count
    expanded: list[str] = []
    size_groups = split_multivalue(size_value)
    for index, part in enumerate(parts):
        count = 1
        if index < len(size_groups):
            match = re.fullmatch(r".+?\*(\d+)\s*个", size_groups[index])
            if match:
                count = int(match.group(1))
        expanded.extend([part] * count)
    if len(expanded) < box_count:
        expanded.extend([expanded[-1]] * (box_count - len(expanded)))
    return expanded[:box_count]


def weight_number(value: str) -> Decimal:
    match = re.search(r"\d+(?:\.\d+)?", clean(value))
    return Decimal(match.group(0)) if match else Decimal("0")


def strip_weight_unit(value: str) -> str:
    return re.sub(r"(?i)\s*kg\s*$", "", clean(value)).strip()


def format_decimal(value: Decimal) -> str:
    if value == value.to_integral():
        return str(int(value))
    return str(value.normalize())


def split_name_spec(text: str) -> tuple[str, str, str]:
    parts = [part.strip() for part in clean(text).split("|")]
    name = parts[0] if parts else clean(text)
    spec = parts[1] if len(parts) > 1 else ""
    material = parts[2] if len(parts) > 2 else ""
    return name, spec, material


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text: str) -> None:
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], text)
        for paragraph in cell.paragraphs[1:]:
            set_paragraph_text(paragraph, "")
    else:
        cell.text = text


def set_cell_lines(cell, lines: list[str]) -> None:
    for index, line in enumerate(lines):
        if index < len(cell.paragraphs):
            set_paragraph_text(cell.paragraphs[index], line)
        else:
            paragraph = cell.add_paragraph()
            set_paragraph_text(paragraph, line)
    for paragraph in cell.paragraphs[len(lines):]:
        set_paragraph_text(paragraph, "")


def add_rows_like_last(table, target_total_rows: int) -> None:
    while len(table.rows) < target_total_rows:
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    while len(table.rows) > target_total_rows:
        table._tbl.remove(table.rows[-1]._tr)


def body_clear(doc: Document) -> None:
    body = doc._body._element
    children = list(body)
    sect_pr = None
    for child in children:
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def append_before_section(doc: Document, element) -> None:
    body = doc._body._element
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(element)
    else:
        body.insert(body.index(sect_pr), element)


def add_page_break_before(paragraph_element) -> None:
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph_element.insert(0, p_pr)
    if p_pr.find(qn("w:pageBreakBefore")) is None:
        p_pr.append(OxmlElement("w:pageBreakBefore"))


def page_break_paragraph() -> OxmlElement:
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run.append(page_break)
    paragraph.append(run)
    return paragraph


def first_packing_form_table(table) -> OxmlElement:
    table_element = deepcopy(table._tbl)
    rows = list(table_element.findall(qn("w:tr")))
    for row in rows[4:]:
        table_element.remove(row)
    return table_element


def adjust_packing_form_rows(table, item_count: int) -> None:
    item_count = max(1, item_count)
    item_template = deepcopy(table.rows[2]._tr)
    while len(table.rows) < item_count + 3:
        table._tbl.insert(table._tbl.index(table.rows[-1]._tr), deepcopy(item_template))
    while len(table.rows) > item_count + 3:
        table._tbl.remove(table.rows[-2]._tr)


def keep_only_first_table(doc: Document) -> None:
    body = doc._body._element
    first_table = doc.tables[0]._element
    for child in list(body):
        if child is first_table or child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def compact_guohe_packing_table(table, font_size: float = 8.5) -> None:
    for base in range(0, len(table.rows), 4):
        header = table.rows[base].cells[0]
        for paragraph_index, paragraph in enumerate(header.paragraphs):
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                if run.text and paragraph_index > 0:
                    run.font.size = Pt(font_size)
        for row_index in range(base + 1, min(base + 4, len(table.rows))):
            for cell in table.rows[row_index].cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        if run.text:
                            run.font.size = Pt(font_size)


def tighten_doc(doc: Document, font_size: float = 9) -> None:
    for section in doc.sections:
        section.top_margin = Pt(20)
        section.bottom_margin = Pt(20)
        section.left_margin = Pt(20)
        section.right_margin = Pt(20)
        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)


def create_info_table(folder: Path, fixed_rows: list[tuple[str, str, str]], batch_rows: list[tuple[str, str, str]]) -> None:
    path = folder / "发货信息搜集表.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "合同固定信息"
    ws.append(["字段", "值", "来源/说明"])
    for row in fixed_rows:
        ws.append(list(row))
    bs = wb.create_sheet("批次发货信息")
    bs.append(["字段", "值", "来源/说明"])
    for row in batch_rows:
        bs.append(list(row))
    hs = wb.create_sheet("填写说明与选项")
    hs.append(["字段", "说明"])
    hs.append(["待确认字段", "计划发运时间、预计到达时间、承运商/车号、收货联系人等未由资料明确给出时，生成待确认版。"])
    fill = PatternFill("solid", fgColor="1F4E78")
    font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin", color="D9E2F3")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = border
                if cell.row == 1:
                    cell.fill = fill
                    cell.font = font
        for col in ("A", "B", "C"):
            sheet.column_dimensions[col].width = 34 if col != "B" else 56
    wb.save(path)


def load_ningde_rows(xlsx_path: Path) -> tuple[list[dict], OrderedDict[int, list[dict]]]:
    wb = load_workbook(xlsx_path, data_only=True)
    ws = wb.active
    source_rows: list[dict] = []
    boxes: OrderedDict[int, list[dict]] = OrderedDict()
    for row in range(2, ws.max_row + 1):
        name = clean(ws.cell(row, 2).value)
        if not name:
            continue
        box_numbers = parse_boxes(ws.cell(row, 6).value)
        size_value = ws.cell(row, 7).value
        weight_value = ws.cell(row, 8).value
        size_raw = clean(size_value)
        sizes = split_counted_parts(size_value, len(box_numbers))
        weights = split_weight_parts(weight_value, size_value, len(box_numbers))
        item = {
            "seq": clean(ws.cell(row, 1).value),
            "name": name,
            "spec": clean(ws.cell(row, 3).value),
            "qty": clean(ws.cell(row, 4).value),
            "unit": clean(ws.cell(row, 5).value),
            "box_raw": clean(ws.cell(row, 6).value),
            "size_raw": size_raw,
            "weight_raw": clean(ws.cell(row, 8).value),
            "pack": clean(ws.cell(row, 9).value),
        }
        source_rows.append(item)
        for index, box in enumerate(box_numbers):
            box_item = dict(item)
            box_item["box"] = box
            box_item["size"] = sizes[index]
            box_item["weight"] = weights[index]
            boxes.setdefault(box, []).append(box_item)
    return source_rows, OrderedDict(sorted(boxes.items()))


def write_guanghe_notice(template: Path, output: Path, source_rows: list[dict], boxes: OrderedDict[int, list[dict]]) -> None:
    doc = Document(template)
    order_no = "20250903（待确认）"
    buyer = "福建宁德核电有限公司"
    project_name = "宁德核电技能培训室维护物项"
    total_boxes = len(boxes)
    total_weight = sum(weight_number(items[0].get("weight", "")) for items in boxes.values())
    heaviest_box, heaviest_items = max(boxes.items(), key=lambda pair: weight_number(pair[1][0].get("weight", "")))
    heaviest = heaviest_items[0]

    replacements = {
        1: f"收货方:{buyer}",
        2: f"发货方:{SUPPLIER}",
        4: f"订单编号:{order_no}",
        5: "2.发货日期：待确认",
        6: "3.到货时间：待确认",
        7: f"4.到货内容：{project_name}",
        8: f"             总件数: {total_boxes}件",
        9: f"             货物总重量：{format_decimal(total_weight)}KG",
        10: f"最重货物尺寸及重量:（{heaviest['size']}）cm，{heaviest['weight']}",
        11: "是否危险品:是（  ）否（√）",
        12: "是否化学品:是（  ）否（√）",
        13: "有无放射性:有（  ）无（√）",
        15: f"承运商/供货商：{SUPPLIER}",
        16: "司机姓名：待确认",
        17: "联系电话：待确认",
        18: "车牌号：待确认",
        20: "7.卸货地点:待确认（宁德核电现场）",
        21: "8.买方接货联系人:待确认",
        22: "买方仓库联系人:待确认",
        35: f"致 : {buyer}",
        38: f"兹证明订单编号{order_no}所提供的物品经检验测试，各项参数符合设计要求及制造标准，同时满足合同中明确的各项技术规范要求。",
        47: f"公司名称：{SUPPLIER}",
        50: "签署人签字：待确认",
    }
    for index, text in replacements.items():
        if index < len(doc.paragraphs):
            set_paragraph_text(doc.paragraphs[index], text)

    pack_table = doc.tables[0]
    add_rows_like_last(pack_table, 6 + len(boxes))
    for row_index in range(6, len(pack_table.rows)):
        for cell in pack_table.rows[row_index].cells:
            set_cell_text(cell, "")
    for row_index, (box, items) in enumerate(boxes.items(), start=6):
        item = items[0]
        values = [f"{box:03d}", f"{item['size']}cm", item["pack"], item["weight"], ""]
        for col, value in enumerate(values):
            set_cell_text(pack_table.rows[row_index].cells[col], value)
    for col in range(1, 5):
        set_cell_text(pack_table.rows[0].cells[col], buyer)
        set_cell_text(pack_table.rows[1].cells[col], SUPPLIER)
        set_cell_text(pack_table.rows[2].cells[col], order_no)
        set_cell_text(pack_table.rows[3].cells[col], "待确认")
        set_cell_text(pack_table.rows[4].cells[col], f"包装箱件数：共 {total_boxes} 件，详情如下：")

    detail_table = doc.tables[1]
    add_rows_like_last(detail_table, 1 + len(source_rows))
    for index, item in enumerate(source_rows, start=1):
        values = [
            str(index * 10),
            item["seq"],
            "NA",
            f"{item['name']} {item['spec']}".strip(),
            item["qty"],
            item["unit"],
            "C3",
            "否",
            "/",
            "/",
        ]
        for col, value in enumerate(values):
            set_cell_text(detail_table.rows[index].cells[col], value)
    tighten_doc(doc, 8.5)
    doc.save(output)


def write_guanghe_marks(template: Path, output: Path, boxes: OrderedDict[int, list[dict]]) -> None:
    doc = Document(template)
    table_template = deepcopy(doc.tables[0]._tbl)
    if len(doc.tables) < len(boxes):
        for _ in range(len(boxes) - len(doc.tables)):
            append_before_section(doc, deepcopy(table_template))
    while len(doc.tables) > len(boxes):
        doc.tables[-1]._element.getparent().remove(doc.tables[-1]._element)
    for table, (box, items) in zip(doc.tables, boxes.items()):
        item = items[0]
        names = []
        for sub in items:
            if sub["name"] not in names:
                names.append(sub["name"])
        device = names[0] if len(names) == 1 else f"{names[0]}等"
        values = {
            (0, 1): f"{box:03d}",
            (0, 5): item["size"],
            (1, 1): strip_weight_unit(item["weight"]),
            (1, 5): strip_weight_unit(item["weight"]),
            (2, 1): device,
            (3, 1): SUPPLIER,
        }
        for (row, col), value in values.items():
            set_cell_text(table.rows[row].cells[col], value)
    tighten_doc(doc, 10)
    doc.save(output)


def load_guohe_rows(xlsx_path: Path) -> tuple[list[dict], OrderedDict[int, list[dict]]]:
    wb = load_workbook(xlsx_path, data_only=True)
    ws = wb.active
    rows: list[dict] = []
    boxes: OrderedDict[int, list[dict]] = OrderedDict()
    for row in range(2, ws.max_row + 1):
        desc = clean(ws.cell(row, 3).value)
        if not desc:
            continue
        name, spec, material = split_name_spec(desc)
        box_numbers = parse_boxes(ws.cell(row, 5).value)
        size_value = ws.cell(row, 7).value
        weight_value = ws.cell(row, 8).value
        size_raw = clean(size_value)
        sizes = split_counted_parts(size_value, len(box_numbers))
        weights = split_weight_parts(weight_value, size_value, len(box_numbers))
        qty = Decimal(clean(ws.cell(row, 4).value) or "0")
        per_box_qty = qty / Decimal(len(box_numbers)) if box_numbers else qty
        item = {
            "line": clean(ws.cell(row, 1).value),
            "code": clean(ws.cell(row, 2).value),
            "desc": desc,
            "name": name,
            "spec": spec,
            "material": material,
            "qty": clean(ws.cell(row, 4).value),
            "unit": "EA",
            "box_raw": clean(ws.cell(row, 5).value),
            "size_raw": size_raw,
            "weight_raw": clean(ws.cell(row, 8).value),
            "pack": clean(ws.cell(row, 9).value),
            "per_box_qty": format_decimal(per_box_qty),
        }
        rows.append(item)
        for index, box in enumerate(box_numbers):
            box_item = dict(item)
            box_item["box"] = box
            box_item["size"] = sizes[index]
            box_item["weight"] = weights[index]
            boxes.setdefault(box, []).append(box_item)
    return rows, OrderedDict(sorted(boxes.items()))


def write_guohe_notice(template: Path, output: Path, source_rows: list[dict], boxes: OrderedDict[int, list[dict]]) -> None:
    doc = Document(template)
    buyer = "国核示范电站有限责任公司"
    order_no = "100601WZ0120250498"
    project_name = "2025年第一批防异物物资采购"
    p_text = {
        1: f"致{buyer}：",
        2: f"我司与贵司签订的订单编号为{order_no}、合同名称为{project_name}，根据合同要求，安排发货，计划发运时间为 待确认  预计到达时间为 待确认 ，贵司收货人为 待确认。",
        4: f" 本次到货批次号为 1 批，共计  {len(source_rows)}   项物资， 是□ 否☑涵盖全部合同物项。",
        10: f"联系人/联系电话：待确认/{SUPPLIER_PHONE}",
    }
    for index, text in p_text.items():
        set_paragraph_text(doc.paragraphs[index], text)

    detail = doc.tables[0]
    add_rows_like_last(detail, 1 + len(source_rows))
    for index, item in enumerate(source_rows, start=1):
        values = [
            str(index),
            item["line"],
            item["code"],
            item["desc"],
            item["unit"],
            item["qty"],
            item["box_raw"],
            "C",
            "有□无☑",
            "无",
            "有",
            "合格证",
        ]
        for col, value in enumerate(values):
            set_cell_text(detail.rows[index].cells[col], value)

    diff = doc.tables[1]
    for col in range(7):
        set_cell_text(diff.rows[1].cells[col], "/")

    pack = doc.tables[2]
    add_rows_like_last(pack, 1 + len(boxes))
    for index, (box, items) in enumerate(boxes.items(), start=1):
        item = items[0]
        values = [str(index), str(box), item["name"], item["size"], item["weight"], item["pack"], "待确认"]
        for col, value in enumerate(values):
            set_cell_text(pack.rows[index].cells[col], value)
    tighten_doc(doc, 7.5)
    doc.save(output)


def write_guohe_marks(template: Path, output: Path, boxes: OrderedDict[int, list[dict]]) -> None:
    template_doc = Document(template)
    title = deepcopy(template_doc.paragraphs[0]._p)
    receiver_table = deepcopy(template_doc.tables[0]._tbl)
    material_table = deepcopy(template_doc.tables[1]._tbl)
    doc = Document(template)
    body_clear(doc)
    for index, (box, items) in enumerate(boxes.items()):
        title_element = deepcopy(title)
        if index:
            add_page_break_before(title_element)
        append_before_section(doc, title_element)
        append_before_section(doc, deepcopy(receiver_table))
        append_before_section(doc, deepcopy(material_table))
    for table_index, (box, items) in enumerate(boxes.items()):
        receiver = doc.tables[table_index * 2]
        material = doc.tables[table_index * 2 + 1]
        item = items[0]
        for row in range(6):
            set_cell_text(receiver.rows[row].cells[0], "国核示范工程")
        set_cell_text(receiver.rows[0].cells[2], "公司名称 ：国核示范电站有限责任公司")
        set_cell_text(receiver.rows[1].cells[2], "收货地址：待确认")
        set_cell_text(receiver.rows[2].cells[2], "联系人：待确认")
        set_cell_text(receiver.rows[2].cells[3], "电话：待确认")
        set_cell_text(receiver.rows[3].cells[2], f"公司名称：{SUPPLIER}")
        set_cell_text(receiver.rows[4].cells[2], f"公司地址：{SUPPLIER_ADDRESS}")
        set_cell_text(receiver.rows[5].cells[2], "联系人：待确认")
        set_cell_text(receiver.rows[5].cells[3], f"电话：{SUPPLIER_PHONE}")
        set_cell_text(material.rows[0].cells[1], "备品备件  ☑")
        set_cell_text(material.rows[1].cells[1], "工器具    □")
        set_cell_text(material.rows[2].cells[1], "耗材      □")
        set_cell_text(material.rows[3].cells[1], "C")
        set_cell_text(material.rows[4].cells[1], item["name"])
    tighten_doc(doc, 9)
    doc.save(output)


def write_guohe_packing_list(template: Path, output: Path, source_rows: list[dict], boxes: OrderedDict[int, list[dict]]) -> None:
    template_doc = Document(template)
    row_templates = [deepcopy(template_doc.tables[0].rows[i]._tr) for i in range(4)]
    doc = Document(template)
    keep_only_first_table(doc)
    table = doc.tables[0]
    while len(table.rows):
        table._tbl.remove(table.rows[0]._tr)
    for box, items in boxes.items():
        base = len(table.rows)
        for row_template in row_templates:
            table._tbl.append(deepcopy(row_template))
        while len(items) > (len(table.rows) - base - 3):
            table._tbl.insert(table._tbl.index(table.rows[base + len(table.rows) - base - 1]._tr), deepcopy(row_templates[2]))
        first_item = items[0]
        header_lines = [
            "装箱单",
            f"箱号：{box}      批次号：1    本批次总共箱件数  {len(boxes)}       储存级别：C",
            f"箱名称：{first_item['name']}         机组号：           ",
            "物资简述：                  ",
            f"箱尺寸（cm）：{first_item['size']}       重量（KG）：  {strip_weight_unit(first_item['weight'])}KG   ",
            "清洁度等级：                  质保等级：  QNA",
            "合同号： 100601WZ0120250498 合同名称： 2025年第一批防异物物资采购    ",
            "币种：RMB              金额：         ",
            f"供应商名称： {SUPPLIER}   供应商地址： {SUPPLIER_ADDRESS}       ",
            f"联系人： 待确认      联系电话：  {SUPPLIER_PHONE}           E_mail：            ",
            "装箱日期： 待确认            ",
        ]
        set_cell_lines(table.rows[base].cells[0], header_lines)
        for row_offset, item in enumerate(items, start=2):
            values = [
                item["line"],
                item["code"],
                item["name"],
                item["spec"],
                "",
                item["per_box_qty"],
                item["per_box_qty"],
                "",
                "",
                "",
                "",
                "待确认",
                item["material"],
                "",
                item["unit"],
                "",
            ]
            for col, value in enumerate(values):
                set_cell_text(table.rows[base + row_offset].cells[col], value)
    compact_guohe_packing_table(table, 8.5)
    doc.save(output)


def write_summary(path: Path, lines: list[str]) -> None:
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def generate_ningde(root: Path, template_root: Path) -> dict:
    folder = root / "宁德核电技能培训室维护物项"
    xlsx = folder / "宁德.xlsx"
    output = folder / "发货资料-第一批-待确认"
    output.mkdir(exist_ok=True)
    source_rows, boxes = load_ningde_rows(xlsx)
    create_info_table(
        folder,
        [
            ("客户体系", "广核", "按宁德核电及模板体系判断"),
            ("项目/合同名称", "宁德核电技能培训室维护物项", "文件夹和 PDF 文件名"),
            ("订单编号", "20250903（待确认）", "PDF 文件名前缀，需用户确认"),
            ("收货方", "福建宁德核电有限公司", "项目名称推断，需用户确认"),
            ("发货方", SUPPLIER, "默认发货单位"),
            ("卸货地点", "待确认（宁德核电现场）", "PDF 无文字层"),
        ],
        [
            ("批次", "第一批", "默认"),
            ("发货清单文件", xlsx.name, "仓库清单"),
            ("计划发运时间", "", "需用户确认"),
            ("预计到达时间", "", "需用户确认"),
            ("承运商/司机/车牌", "", "需用户确认"),
        ],
    )
    write_guanghe_notice(template_root / "广核" / "国内到货通知.docx", output / "国内到货通知-第一批-待确认.docx", source_rows, boxes)
    write_guanghe_marks(template_root / "广核" / "唛头.docx", output / "唛头-第一批-待确认.docx", boxes)
    write_summary(
        output / "生成说明-第一批-待确认.md",
        [
            "# 生成说明",
            "",
            "- 客户体系：广核（按宁德核电项目判断）。",
            f"- 物项明细：{len(source_rows)} 项；箱号：{min(boxes)}-{max(boxes)}，共 {len(boxes)} 件。",
            "- 已生成：国内到货通知、唛头。",
            "- 宁德 PDF 无可抽取文字层，订单编号、收货联系人、卸货地点等需用户核对。",
            "- 待确认字段：发货日期、到货时间、承运商、司机、联系电话、车牌号、买方接货联系人、仓库联系人、签署人。",
        ],
    )
    return {"folder": str(output), "items": len(source_rows), "boxes": len(boxes)}


def generate_guohe(root: Path, template_root: Path) -> dict:
    folder = root / "2025年第一批防异物物资采购"
    xlsx = next(folder.glob("*.xlsx"))
    output = folder / "发货资料-第一批-待确认"
    output.mkdir(exist_ok=True)
    source_rows, boxes = load_guohe_rows(xlsx)
    create_info_table(
        folder,
        [
            ("客户体系", "国核", "文件名含国核示范订单"),
            ("项目/合同名称", "2025年第一批防异物物资采购", "PDF 文件名"),
            ("订单编号", "100601WZ0120250498", "Excel 文件名"),
            ("收货方", "国核示范电站有限责任公司", "按国核示范订单判断，需用户确认"),
            ("发货方", SUPPLIER, "默认发货单位"),
            ("收货地址/联系人", "待确认", "PDF 提取文字乱码，需用户确认"),
        ],
        [
            ("批次", "第一批", "文件名"),
            ("发货清单文件", xlsx.name, "仓库清单"),
            ("计划发运时间", "", "需用户确认"),
            ("预计到达时间", "", "需用户确认"),
            ("卸货方式", "", "需用户确认"),
        ],
    )
    write_guohe_notice(template_root / "国核" / "到 货 通 知 单（第一批）.docx", output / "到货通知单-第一批-待确认.docx", source_rows, boxes)
    write_guohe_marks(template_root / "国核" / "唛头.docx", output / "唛头-第一批-待确认.docx", boxes)
    write_guohe_packing_list(template_root / "国核" / "装箱单（第一批）.docx", output / "装箱单-第一批-待确认.docx", source_rows, boxes)
    write_summary(
        output / "生成说明-第一批-待确认.md",
        [
            "# 生成说明",
            "",
            "- 客户体系：国核。",
            f"- 物项明细：{len(source_rows)} 项；箱号：{min(boxes)}-{max(boxes)}，共 {len(boxes)} 件。",
            "- 已生成：到货通知单、唛头、装箱单。",
            "- 箱号范围已展开；装箱单中按箱号平均拆分数量，需仓库核对确认。",
            "- 待确认字段：计划发运时间、预计到达时间、收货地址、收货联系人、发货联系人、卸货方式、制造商、装箱日期。",
        ],
    )
    return {"folder": str(output), "items": len(source_rows), "boxes": len(boxes)}


def main() -> int:
    parser = argparse.ArgumentParser(description="生成广核/国核发货资料待确认版。")
    parser.add_argument("root", help="发货通知工作根目录。")
    parser.add_argument("--templates", required=True, help="发货通知模板目录。")
    parser.add_argument("--project", choices=["ningde", "guohe", "all"], default="all")
    args = parser.parse_args()

    root = Path(args.root).resolve()
    template_root = Path(args.templates).resolve()
    results = {}
    if args.project in {"ningde", "all"}:
        results["ningde"] = generate_ningde(root, template_root)
    if args.project in {"guohe", "all"}:
        results["guohe"] = generate_guohe(root, template_root)
    print(json.dumps(results, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
