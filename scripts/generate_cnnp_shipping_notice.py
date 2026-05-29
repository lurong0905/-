#!/usr/bin/env python
# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import json
import re
import shutil
from collections import defaultdict
from copy import deepcopy
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


INFO_TABLE_NAME = "发货信息搜集表.xlsx"
DEFAULT_OUTPUT_FOLDER = "发货资料-{batch}-新流程-待确认"


def clean(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def read_info_table(info_path: Path) -> tuple[dict[str, str], dict[str, str]]:
    wb = load_workbook(info_path, data_only=True)
    fixed_ws = wb["合同固定信息"]
    batch_ws = wb["批次发货信息"]

    fixed: dict[str, str] = {}
    for row in fixed_ws.iter_rows(min_row=2, values_only=True):
        if row and row[0]:
            fixed[clean(row[0])] = clean(row[1])

    headers = [clean(cell.value) for cell in batch_ws[1]]
    batch: dict[str, str] = {}
    for row in batch_ws.iter_rows(min_row=2, values_only=True):
        if row and any(value is not None for value in row):
            batch = {
                headers[index]: clean(row[index]) if index < len(row) else ""
                for index in range(len(headers))
            }
            break
    return fixed, batch


def merged_value_reader(ws):
    merged = {}
    for rng in ws.merged_cells.ranges:
        value = ws.cell(rng.min_row, rng.min_col).value
        for row in range(rng.min_row, rng.max_row + 1):
            for col in range(rng.min_col, rng.max_col + 1):
                merged[(row, col)] = value

    def cell(row: int, col: int):
        return merged.get((row, col), ws.cell(row, col).value)

    return cell


def parse_boxes(value: str) -> list[int]:
    text = clean(value)
    if not text:
        return []
    match = re.fullmatch(r"(\d+)\s*[-~至]\s*(\d+)", text)
    if match:
        start, end = int(match.group(1)), int(match.group(2))
        return list(range(start, end + 1))
    return [int(item) for item in re.findall(r"\d+", text)]


def item_base_name(name: str) -> str:
    text = clean(name).replace("\r", "\n")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\b([A-Za-z])\s+(?=[\u4e00-\u9fff])", r"\1", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"^[A-Za-z0-9_-]+\s*[/／\\]\s*", "", text).strip()
    text = re.split(r"[/／\\]", text, maxsplit=1)[0].strip()
    parts = text.split()
    if parts and re.search(r"[\u4e00-\u9fff]", parts[0]):
        return parts[0]
    match = re.search(r"[A-Za-z]?[\u4e00-\u9fff][A-Za-z\u4e00-\u9fff（）()、·-]*", text)
    return match.group(0) if match else text


def read_shipment_rows(shipment_path: Path) -> list[dict[str, str]]:
    wb = load_workbook(shipment_path, data_only=True)
    ws = wb.active
    cell = merged_value_reader(ws)
    rows: list[dict[str, str]] = []
    for row in range(2, ws.max_row + 1):
        item = {
            "合同序号": clean(cell(row, 1)),
            "田湾编码": clean(cell(row, 2)),
            "物项编码": clean(cell(row, 3)),
            "物项名称": clean(cell(row, 4)),
            "数量": clean(cell(row, 5)),
            "单位": clean(cell(row, 6)),
            "箱号": clean(cell(row, 7)),
            "包装尺寸": clean(cell(row, 8)),
            "毛重/净重": clean(cell(row, 9)),
            "包装": clean(cell(row, 10)),
            "备注": clean(cell(row, 11)),
        }
        if any(item.values()):
            rows.append(item)
    return rows


def summarize_shipment(items: list[dict[str, str]]) -> dict[str, object]:
    box_set: set[int] = set()
    box_items: dict[int, set[str]] = defaultdict(set)
    box_pack: dict[int, str] = {}
    for item in items:
        for box in parse_boxes(item["箱号"]):
            box_set.add(box)
            box_items[box].add(item["物项编码"] or item["物项名称"] or item["合同序号"])
            if item["包装"]:
                box_pack[box] = item["包装"]

    pack_counts: dict[str, int] = defaultdict(int)
    for box in box_set:
        pack_counts[box_pack.get(box, "未注明")] += 1

    unique_names: list[str] = []
    for item in items:
        name = item_base_name(item["物项名称"])
        if name and name not in unique_names:
            unique_names.append(name)

    main_desc = f"{'、'.join(unique_names[:5])}等。" if unique_names else ""

    pack_desc = "；".join(f"{name}：{count}件" for name, count in sorted(pack_counts.items()))
    if pack_desc:
        pack_desc += f"；合计：{len(box_set)}箱/件"

    qualified_boxes = [box for box, values in sorted(box_items.items()) if len(values) >= 3]
    return {
        "box_count": len(box_set),
        "main_desc": main_desc,
        "pack_desc": pack_desc,
        "qualified_boxes": qualified_boxes,
    }


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell_obj, text: str) -> None:
    if cell_obj.paragraphs:
        set_paragraph_text(cell_obj.paragraphs[0], text)
        for paragraph in cell_obj.paragraphs[1:]:
            set_paragraph_text(paragraph, "")
    else:
        cell_obj.text = text


def fill_cell(table, row: int, col: int, text: str) -> None:
    set_cell_text(table.rows[row].cells[col], text)


def add_table_rows_like_last(table, target_data_rows: int) -> None:
    while len(table.rows) - 1 < target_data_rows:
        table._tbl.append(deepcopy(table.rows[-1]._tr))


def render_notice(template_path: Path, output_path: Path, fixed: dict[str, str], batch: dict[str, str], items: list[dict[str, str]]) -> dict[str, object]:
    summary = summarize_shipment(items)
    doc = Document(template_path)

    batch_name = batch.get("批次") or "第一批"
    plan_ship_time = batch.get("计划发运时间")
    ship_date_text = plan_ship_time or "____年__月__日"
    arrival_date = batch.get("预计到达时间")
    waybill = batch.get("运单号/快递单号")
    transport = batch.get("运输方式")
    unload = batch.get("装卸货方式")
    storage_level = batch.get("储存级别")
    storage_condition = batch.get("储存条件")

    contract_name = fixed.get("合同名称")
    contract_no = fixed.get("合同号")
    buyer = fixed.get("买方名称")
    seller = fixed.get("卖方名称")
    receiver = "，".join(part for part in [fixed.get("收货人"), fixed.get("收货电话")] if part)
    contact_manager = fixed.get("合同管理联系人")
    delivery_requirement = fixed.get("合同要求交货时间") or "根据需求供货"
    ship_from = fixed.get("发货地点")
    sender = fixed.get("发货单位") or seller
    sender_contact = fixed.get("发货单位联系人")
    sender_phone = fixed.get("发货单位联系电话")
    project_no = fixed.get("立项单号默认值") or "/"
    project_owner = fixed.get("采购申请处室项目负责人默认值") or "/"

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.strip() == "发货通知（第X批）":
            set_paragraph_text(paragraph, f"发货通知（{batch_name}）")
        elif text.strip() == "买方名称：":
            set_paragraph_text(paragraph, f"买方名称：{buyer}")
        elif "我公司将于xxxx年x月xx日发运" in text:
            set_paragraph_text(paragraph, f"按照采购合同规定，我公司将于{ship_date_text}发运，请查收。")
        elif "根据《合同名称》相关规定" in text:
            set_paragraph_text(
                paragraph,
                f"根据《{contract_name}》相关规定，我司将于{ship_date_text}发运合同物项，发运前我司已经对拟发货物项进行了自检，现将自检结果报告贵司。",
            )
        elif text.strip() == "Xxxx年x月xx日":
            set_paragraph_text(paragraph, ship_date_text)

    main_table = doc.tables[0]
    fill_cell(main_table, 0, 1, contract_name)
    fill_cell(main_table, 0, 5, contract_no)
    fill_cell(main_table, 1, 1, delivery_requirement)
    fill_cell(main_table, 2, 1, plan_ship_time or "")
    fill_cell(main_table, 2, 4, arrival_date)
    fill_cell(main_table, 3, 1, receiver)
    fill_cell(main_table, 3, 4, contact_manager)
    fill_cell(main_table, 4, 1, project_no)
    fill_cell(main_table, 4, 4, project_owner)
    fill_cell(main_table, 5, 1, ship_from)
    fill_cell(main_table, 6, 1, str(summary["main_desc"]))
    fill_cell(main_table, 7, 1, waybill)
    fill_cell(main_table, 8, 1, "非寿期物项 □寿期物项（已在装箱清单中注明寿期）")
    fill_cell(main_table, 9, 1, str(summary["pack_desc"]))
    fill_cell(main_table, 10, 1, f"运输方式：{transport or '          '}    车号：{waybill or '          '}")
    if unload:
        fill_cell(main_table, 11, 1, unload)
    if storage_level or storage_condition:
        fill_cell(main_table, 12, 1, f"储存级别 {storage_level} 级   {storage_condition}")
    fill_cell(main_table, 13, 1, sender)
    fill_cell(main_table, 14, 1, sender_contact)
    fill_cell(main_table, 14, 5, sender_phone)
    fill_cell(main_table, 15, 1, "1、装箱清单  2、随箱文件清单  3、合同物项自检报告")

    packing_table = doc.tables[1]
    add_table_rows_like_last(packing_table, len(items))
    for index, item in enumerate(items, start=1):
        values = [
            item["合同序号"],
            item["田湾编码"],
            item["物项编码"],
            item["物项名称"],
            item["数量"],
            item["单位"],
            item["箱号"],
            item["包装尺寸"],
            item["毛重/净重"],
            item["备注"],
        ]
        for col, value in enumerate(values):
            fill_cell(packing_table, index, col, value)

    self_check = doc.tables[2]
    fill_cell(self_check, 6, 3, "装箱清单、随箱文件清单、合同物项自检报告")

    doc.save(output_path)

    missing = [
        label
        for label, value in [
            ("计划发运时间", plan_ship_time),
            ("预计到达时间", arrival_date),
            ("装卸货方式", unload),
            ("储存级别", storage_level),
            ("储存条件", storage_condition),
            ("发货单位联系人", sender_contact),
        ]
        if not value
    ]
    return {
        "output": str(output_path),
        "batch": batch_name,
        "items": len(items),
        "box_count": summary["box_count"],
        "qualified_packing_list_boxes": summary["qualified_boxes"],
        "missing_fields": missing,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="按中核模板生成发货通知待确认版。")
    parser.add_argument("contract_folder", help="合同文件夹路径。")
    parser.add_argument("--template", required=True, help="中核发货通知模板 docx 路径。")
    parser.add_argument("--output-folder", help="输出文件夹路径。")
    args = parser.parse_args()

    contract_folder = Path(args.contract_folder).resolve()
    template_path = Path(args.template).resolve()
    info_path = contract_folder / INFO_TABLE_NAME
    if not info_path.exists():
        raise SystemExit(f"未找到信息搜集表：{info_path}")
    if not template_path.exists():
        raise SystemExit(f"未找到发货通知模板：{template_path}")

    fixed, batch = read_info_table(info_path)
    batch_name = batch.get("批次") or "第一批"
    shipment_name = batch.get("发货清单文件") or "发货清单（第一批）.xlsx"
    shipment_path = contract_folder / shipment_name
    if not shipment_path.exists():
        raise SystemExit(f"未找到发货清单：{shipment_path}")

    output_folder = Path(args.output_folder).resolve() if args.output_folder else contract_folder / DEFAULT_OUTPUT_FOLDER.format(batch=batch_name)
    output_folder.mkdir(parents=True, exist_ok=True)
    output_path = output_folder / f"发货通知-{batch_name}-待确认.docx"

    shutil.copy2(template_path, output_path)
    result = render_notice(output_path, output_path, fixed, batch, read_shipment_rows(shipment_path))
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
