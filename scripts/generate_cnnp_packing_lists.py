#!/usr/bin/env python
# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import json
import re
from collections import OrderedDict
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from openpyxl import load_workbook


INFO_TABLE_NAME = "发货信息搜集表.xlsx"
DEFAULT_OUTPUT_FOLDER = "发货资料-{batch}-新流程-待确认"
TITLE_FONT_SIZE_PT = 16
TABLE_FONT_SIZE_PT = 8
HEADER_ROW_HEIGHT_PT = 20
DATA_ROW_HEIGHT_PT = 24


def clean(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def read_info_table(info_path: Path) -> tuple[dict[str, str], dict[str, str]]:
    wb = load_workbook(info_path, data_only=True)
    fixed_ws = wb["合同固定信息"]
    batch_ws = wb["批次发货信息"]

    fixed: dict[str, str] = {}
    for row in fixed_ws.iter_rows(min_row=2, values_only=True):
        if row and row[0]:
            fixed[clean(row[0])] = clean(row[1])

    headers = [clean(cell.value) for cell in batch_ws[1]]
    batch: dict[str, str] = {}
    for row in batch_ws.iter_rows(min_row=2, values_only=True):
        if row and any(value is not None for value in row):
            batch = {
                headers[index]: clean(row[index]) if index < len(row) else ""
                for index in range(len(headers))
            }
            break
    return fixed, batch


def merged_value_reader(ws):
    merged = {}
    for rng in ws.merged_cells.ranges:
        value = ws.cell(rng.min_row, rng.min_col).value
        for row in range(rng.min_row, rng.max_row + 1):
            for col in range(rng.min_col, rng.max_col + 1):
                merged[(row, col)] = value

    def cell(row: int, col: int):
        return merged.get((row, col), ws.cell(row, col).value)

    return cell


def parse_boxes(value: str) -> list[int]:
    text = clean(value)
    if not text:
        return []
    match = re.fullmatch(r"(\d+)\s*[-~至]\s*(\d+)", text)
    if match:
        start, end = int(match.group(1)), int(match.group(2))
        return list(range(start, end + 1))
    return [int(item) for item in re.findall(r"\d+", text)]


def read_shipment_rows(shipment_path: Path) -> list[dict[str, str]]:
    wb = load_workbook(shipment_path, data_only=True)
    ws = wb.active
    cell = merged_value_reader(ws)
    rows: list[dict[str, str]] = []
    last_box = last_size = last_weight = last_pack = ""
    for row in range(2, ws.max_row + 1):
        box = clean(cell(row, 7)) or last_box
        size = clean(cell(row, 8)) or last_size
        weight = clean(cell(row, 9)) or last_weight
        pack = clean(cell(row, 10)) or last_pack
        if box:
            last_box = box
        if size:
            last_size = size
        if weight:
            last_weight = weight
        if pack:
            last_pack = pack
        item = {
            "合同序号": clean(cell(row, 1)),
            "田湾编码": clean(cell(row, 2)),
            "物项编码": clean(cell(row, 3)),
            "物项名称": clean(cell(row, 4)),
            "数量": clean(cell(row, 5)),
            "单位": clean(cell(row, 6)),
            "箱号": box,
            "包装尺寸": size,
            "毛重/净重": weight,
            "包装": pack,
            "备注": clean(cell(row, 11)),
        }
        if any(item.values()):
            rows.append(item)
    return rows


def item_name_and_spec(value: str) -> tuple[str, str]:
    text = clean(value).replace("\r", "\n")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\b([A-Za-z])\s+(?=[\u4e00-\u9fff])", r"\1", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"^[A-Za-z0-9_-]+\s*[/／\\]\s*", "", text).strip()
    text = re.split(r"[/／\\]", text, maxsplit=1)[0].strip()
    parts = text.split(maxsplit=1)
    if parts and re.search(r"[\u4e00-\u9fff]", parts[0]):
        return parts[0], parts[1] if len(parts) > 1 else ""
    match = re.search(r"[A-Za-z]?[\u4e00-\u9fff][A-Za-z\u4e00-\u9fff（）()、·-]*", text)
    if not match:
        return text, ""
    name = match.group(0)
    spec = text[match.end():].strip()
    return name, spec


def group_items_by_box(items: list[dict[str, str]]) -> OrderedDict[int, list[dict[str, str]]]:
    grouped: OrderedDict[int, list[dict[str, str]]] = OrderedDict()
    for item in items:
        for box in parse_boxes(item["箱号"]):
            grouped.setdefault(box, []).append(item)
    return OrderedDict(sorted(grouped.items(), key=lambda pair: pair[0]))


def qualified_boxes(box_groups: OrderedDict[int, list[dict[str, str]]]) -> OrderedDict[int, list[dict[str, str]]]:
    qualified: OrderedDict[int, list[dict[str, str]]] = OrderedDict()
    for box, items in box_groups.items():
        categories = {item["物项编码"] or item["物项名称"] or item["合同序号"] for item in items}
        categories.discard("")
        if len(categories) >= 3:
            qualified[box] = items
    return qualified


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    children = list(body)
    sect_pr = None
    for child in children:
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def append_before_section(doc: Document, element) -> None:
    body = doc._body._element
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(element)
    else:
        body.insert(body.index(sect_pr), element)


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell_obj, text: str) -> None:
    if cell_obj.paragraphs:
        set_paragraph_text(cell_obj.paragraphs[0], text)
        for paragraph in cell_obj.paragraphs[1:]:
            set_paragraph_text(paragraph, "")
    else:
        cell_obj.text = text


def add_table_rows_like_last(table: Table, target_data_rows: int) -> None:
    while len(table.rows) - 1 < target_data_rows:
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    while len(table.rows) - 1 > target_data_rows:
        table._tbl.remove(table.rows[-1]._tr)


def fill_table(table: Table, box_no: int, items: list[dict[str, str]]) -> None:
    add_table_rows_like_last(table, len(items))
    for index, item in enumerate(items, start=1):
        name, spec = item_name_and_spec(item["物项名称"])
        remark = item["备注"] or (f"包装：{item['包装']}" if item["包装"] else "")
        values = [
            item["合同序号"],
            item["物项编码"],
            item["田湾编码"],
            name,
            spec,
            item["单位"],
            item["数量"],
            str(box_no),
            item["包装尺寸"],
            item["毛重/净重"],
            remark,
        ]
        for col, value in enumerate(values):
            set_cell_text(table.rows[index].cells[col], value)


def tighten_layout(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Pt(18)
        section.bottom_margin = Pt(18)
        section.left_margin = Pt(18)
        section.right_margin = Pt(18)
        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            if paragraph.text.strip() == "装箱清单":
                run.font.size = Pt(TITLE_FONT_SIZE_PT)
                run.font.bold = True
    for table in doc.tables:
        table.allow_autofit = True
        for row_index, row in enumerate(table.rows):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Pt(HEADER_ROW_HEIGHT_PT if row_index == 0 else DATA_ROW_HEIGHT_PT)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(TABLE_FONT_SIZE_PT)
                        if row_index == 0:
                            run.font.bold = True


def render_packing_lists(template_path: Path, output_folder: Path, batch_name: str, box_groups: OrderedDict[int, list[dict[str, str]]]) -> dict[str, object]:
    template_doc = Document(template_path)
    title_element = deepcopy(template_doc.paragraphs[0]._p)
    table_element = deepcopy(template_doc.tables[0]._tbl)

    outputs: list[str] = []
    for box_no, items in box_groups.items():
        doc = Document(template_path)
        clear_document_body(doc)
        append_before_section(doc, deepcopy(title_element))
        append_before_section(doc, deepcopy(table_element))
        fill_table(doc.tables[0], box_no, items)
        tighten_layout(doc)
        output_path = output_folder / f"装箱清单（每箱1张）-{box_no}号箱-{batch_name}-待确认.docx"
        doc.save(output_path)
        outputs.append(str(output_path))

    return {
        "outputs": outputs,
        "boxes": list(box_groups.keys()),
        "count": len(outputs),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="按中核模板生成达到规则的每箱装箱清单。")
    parser.add_argument("contract_folder", help="合同文件夹路径。")
    parser.add_argument("--template", required=True, help="中核装箱清单模板 docx 路径。")
    parser.add_argument("--output-folder", help="输出文件夹路径。")
    args = parser.parse_args()

    contract_folder = Path(args.contract_folder).resolve()
    template_path = Path(args.template).resolve()
    info_path = contract_folder / INFO_TABLE_NAME
    if not info_path.exists():
        raise SystemExit(f"未找到信息搜集表：{info_path}")
    if not template_path.exists():
        raise SystemExit(f"未找到装箱清单模板：{template_path}")

    _fixed, batch = read_info_table(info_path)
    batch_name = batch.get("批次") or "第一批"
    shipment_name = batch.get("发货清单文件") or f"发货清单（{batch_name}）.xlsx"
    shipment_path = contract_folder / shipment_name
    if not shipment_path.exists():
        raise SystemExit(f"未找到发货清单：{shipment_path}")

    output_folder = (
        Path(args.output_folder).resolve()
        if args.output_folder
        else contract_folder / DEFAULT_OUTPUT_FOLDER.format(batch=batch_name)
    )
    output_folder.mkdir(parents=True, exist_ok=True)
    all_box_groups = group_items_by_box(read_shipment_rows(shipment_path))
    result = render_packing_lists(template_path, output_folder, batch_name, qualified_boxes(all_box_groups))
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
