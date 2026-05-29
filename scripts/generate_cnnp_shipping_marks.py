#!/usr/bin/env python
# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import json
import re
from collections import OrderedDict
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from openpyxl import load_workbook


INFO_TABLE_NAME = "发货信息搜集表.xlsx"
DEFAULT_OUTPUT_FOLDER = "发货资料-{batch}-新流程-修正版-待确认"
TITLE_FONT_SIZE_PT = 16
LABEL_FONT_SIZE_PT = 8.5
VALUE_FONT_SIZE_PT = 8.5
VALUE_FONT_BOLD = True
TABLE_TOP_TWIPS = 0
MARK_ROW_OFFSET = 1


def clean(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def read_info_table(info_path: Path) -> tuple[dict[str, str], dict[str, str]]:
    wb = load_workbook(info_path, data_only=True)
    fixed_ws = wb["合同固定信息"]
    batch_ws = wb["批次发货信息"]

    fixed: dict[str, str] = {}
    for row in fixed_ws.iter_rows(min_row=2, values_only=True):
        if row and row[0]:
            fixed[clean(row[0])] = clean(row[1])

    headers = [clean(cell.value) for cell in batch_ws[1]]
    batch: dict[str, str] = {}
    for row in batch_ws.iter_rows(min_row=2, values_only=True):
        if row and any(value is not None for value in row):
            batch = {
                headers[index]: clean(row[index]) if index < len(row) else ""
                for index in range(len(headers))
            }
            break
    return fixed, batch


def merged_value_reader(ws):
    merged = {}
    for rng in ws.merged_cells.ranges:
        value = ws.cell(rng.min_row, rng.min_col).value
        for row in range(rng.min_row, rng.max_row + 1):
            for col in range(rng.min_col, rng.max_col + 1):
                merged[(row, col)] = value

    def cell(row: int, col: int):
        return merged.get((row, col), ws.cell(row, col).value)

    return cell


def parse_boxes(value: str) -> list[int]:
    text = clean(value)
    if not text:
        return []
    match = re.fullmatch(r"(\d+)\s*[-~至]\s*(\d+)", text)
    if match:
        start, end = int(match.group(1)), int(match.group(2))
        return list(range(start, end + 1))
    return [int(item) for item in re.findall(r"\d+", text)]


def read_shipment_rows(shipment_path: Path) -> list[dict[str, str]]:
    wb = load_workbook(shipment_path, data_only=True)
    ws = wb.active
    cell = merged_value_reader(ws)
    rows: list[dict[str, str]] = []
    for row in range(2, ws.max_row + 1):
        item = {
            "合同序号": clean(cell(row, 1)),
            "田湾编码": clean(cell(row, 2)),
            "物项编码": clean(cell(row, 3)),
            "物项名称": clean(cell(row, 4)),
            "数量": clean(cell(row, 5)),
            "单位": clean(cell(row, 6)),
            "箱号": clean(cell(row, 7)),
            "包装尺寸": clean(cell(row, 8)),
            "毛重/净重": clean(cell(row, 9)),
            "包装": clean(cell(row, 10)),
            "备注": clean(cell(row, 11)),
        }
        if any(item.values()):
            rows.append(item)
    return rows


def group_items_by_box(items: list[dict[str, str]]) -> tuple[OrderedDict[int, list[dict[str, str]]], int]:
    grouped: OrderedDict[int, list[dict[str, str]]] = OrderedDict()
    for item in items:
        for box in parse_boxes(item["箱号"]):
            grouped.setdefault(box, []).append(item)
    grouped = OrderedDict(sorted(grouped.items(), key=lambda pair: pair[0]))
    max_box = max(grouped.keys()) if grouped else 0
    return grouped, max_box


def unique_join(values: list[str]) -> str:
    seen: list[str] = []
    for value in values:
        if value and value not in seen:
            seen.append(value)
    return "、".join(seen)


def item_base_name(name: str) -> str:
    text = clean(name).replace("\r", "\n")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\b([A-Za-z])\s+(?=[\u4e00-\u9fff])", r"\1", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"^[A-Za-z0-9_-]+\s*[/／\\]\s*", "", text).strip()
    text = re.split(r"[/／\\]", text, maxsplit=1)[0].strip()
    parts = text.split()
    if parts and re.search(r"[\u4e00-\u9fff]", parts[0]):
        return parts[0]
    match = re.search(r"[A-Za-z]?[\u4e00-\u9fff][A-Za-z\u4e00-\u9fff（）()、·-]*", text)
    return match.group(0) if match else text


def box_item_name(items: list[dict[str, str]]) -> str:
    full_names = [item["物项名称"] for item in items if item["物项名称"]]
    base_names = [item_base_name(name) for name in full_names]
    unique_bases: list[str] = []
    for name in base_names:
        if name and name not in unique_bases:
            unique_bases.append(name)
    if not unique_bases:
        return ""
    if len(unique_bases) == 1:
        return unique_bases[0]
    return f"{unique_bases[0]}等"


def box_first_value(items: list[dict[str, str]], key: str) -> str:
    for item in items:
        if item.get(key):
            return item[key]
    return ""


def set_paragraph_text(paragraph, text: str, font_size_pt: float | None = None, bold: bool | None = None) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    if font_size_pt is not None:
        for run in paragraph.runs:
            run.font.size = Pt(font_size_pt)
    if bold is not None:
        for run in paragraph.runs:
            run.font.bold = bold


def set_cell_text(cell_obj, text: str, font_size_pt: float | None = None, bold: bool | None = None) -> None:
    if cell_obj.paragraphs:
        set_paragraph_text(cell_obj.paragraphs[0], text, font_size_pt=font_size_pt, bold=bold)
        for paragraph in cell_obj.paragraphs[1:]:
            set_paragraph_text(paragraph, "", font_size_pt=font_size_pt, bold=bold)
    else:
        cell_obj.text = text


def fill_cell(table: Table, row: int, col: int, text: str, font_size_pt: float | None = None, bold: bool | None = None) -> None:
    set_cell_text(table.rows[row].cells[col], text, font_size_pt=font_size_pt, bold=bold)


def fill_value_cell(table: Table, row: int, col: int, text: str) -> None:
    fill_cell(table, row, col, text, font_size_pt=VALUE_FONT_SIZE_PT, bold=VALUE_FONT_BOLD)


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    children = list(body)
    sect_pr = None
    for child in children:
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def append_title_before_section(doc: Document, add_page_break: bool = False) -> None:
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    if add_page_break:
        p_pr.append(OxmlElement("w:pageBreakBefore"))
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    p_pr.append(jc)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    paragraph.append(p_pr)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_pr.append(OxmlElement("w:b"))
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(TITLE_FONT_SIZE_PT * 2)))
    r_pr.append(size)
    east_asia_size = OxmlElement("w:szCs")
    east_asia_size.set(qn("w:val"), str(int(TITLE_FONT_SIZE_PT * 2)))
    r_pr.append(east_asia_size)
    run.append(r_pr)

    text = OxmlElement("w:t")
    text.text = "唛  头"
    run.append(text)
    paragraph.append(run)

    body = doc._body._element
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(paragraph)
    else:
        body.insert(body.index(sect_pr), paragraph)


def append_table_before_section(doc: Document, table_element) -> None:
    body = doc._body._element
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(table_element)
    else:
        body.insert(body.index(sect_pr), table_element)


def append_paragraph_before_section(doc: Document, paragraph_element) -> None:
    body = doc._body._element
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(paragraph_element)
    else:
        body.insert(body.index(sect_pr), paragraph_element)


def remove_embedded_section_properties(paragraph_element) -> None:
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is not None:
        for sect_pr in list(p_pr.findall(qn("w:sectPr"))):
            p_pr.remove(sect_pr)


def set_page_break_before_table(table: Table) -> None:
    paragraph = table.rows[0].cells[0].paragraphs[0]
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:pageBreakBefore")) is None:
        p_pr.append(OxmlElement("w:pageBreakBefore"))


def set_table_vertical_position(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    tblp_pr = tbl_pr.find(qn("w:tblpPr"))
    if tblp_pr is not None:
        tblp_pr.set(qn("w:tblpY"), str(TABLE_TOP_TWIPS))


def set_table_inline(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    tblp_pr = tbl_pr.find(qn("w:tblpPr"))
    if tblp_pr is not None:
        tbl_pr.remove(tblp_pr)


def set_header_title(doc: Document) -> None:
    for section in doc.sections:
        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        set_paragraph_text(paragraph, "唛  头", font_size_pt=TITLE_FONT_SIZE_PT, bold=True)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1


def tighten_page_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Pt(0)
        section.bottom_margin = Pt(0)
        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)


def apply_table_font(table: Table) -> None:
    value_cells = {
        (MARK_ROW_OFFSET + 0, 1), (MARK_ROW_OFFSET + 1, 1),
        (MARK_ROW_OFFSET + 2, 1), (MARK_ROW_OFFSET + 3, 1),
        (MARK_ROW_OFFSET + 4, 1), (MARK_ROW_OFFSET + 5, 1),
        (MARK_ROW_OFFSET + 6, 1), (MARK_ROW_OFFSET + 7, 1),
        (MARK_ROW_OFFSET + 7, 3), (MARK_ROW_OFFSET + 8, 1),
        (MARK_ROW_OFFSET + 8, 3), (MARK_ROW_OFFSET + 9, 1),
        (MARK_ROW_OFFSET + 9, 3),
    }
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if row_index == 0:
                size = TITLE_FONT_SIZE_PT
            else:
                size = VALUE_FONT_SIZE_PT if (row_index, col_index) in value_cells else LABEL_FONT_SIZE_PT
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size)
                    run.font.bold = True
    set_table_vertical_position(table)


def insert_title_row(table: Table) -> None:
    title_row = OxmlElement("w:tr")
    tr_pr = OxmlElement("w:trPr")
    height = OxmlElement("w:trHeight")
    height.set(qn("w:val"), "220")
    height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(height)
    title_row.append(tr_pr)

    title_cell = OxmlElement("w:tc")
    tc_pr = OxmlElement("w:tcPr")
    grid_span = OxmlElement("w:gridSpan")
    grid_span.set(qn("w:val"), "4")
    tc_pr.append(grid_span)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    tc_pr.append(borders)
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), "center")
    tc_pr.append(v_align)
    title_cell.append(tc_pr)

    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    p_pr.append(jc)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    paragraph.append(p_pr)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_pr.append(OxmlElement("w:b"))
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(TITLE_FONT_SIZE_PT * 2)))
    r_pr.append(size)
    east_asia_size = OxmlElement("w:szCs")
    east_asia_size.set(qn("w:val"), str(int(TITLE_FONT_SIZE_PT * 2)))
    r_pr.append(east_asia_size)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "唛  头"
    run.append(text)
    paragraph.append(run)
    title_cell.append(paragraph)
    title_row.append(title_cell)

    tbl = table._tbl
    first_row_index = next(index for index, child in enumerate(tbl) if child.tag == qn("w:tr"))
    tbl.insert(first_row_index, title_row)


def fill_mark_table(table: Table, fixed: dict[str, str], box_no: int, total_count: int, items: list[dict[str, str]]) -> None:
    buyer = fixed.get("买方名称", "")
    contract_manager = fixed.get("合同管理联系人", "")
    receiver = fixed.get("收货人", "")
    receiver_phone = fixed.get("收货电话", "")
    sender = fixed.get("发货单位") or fixed.get("卖方名称", "")
    ship_from = fixed.get("发货地点", "")
    ship_from_text = "连云港" if "连云港" in ship_from else ship_from
    destination = fixed.get("收货地址", "")

    fill_value_cell(table, MARK_ROW_OFFSET + 0, 1, fixed.get("合同号", ""))
    fill_value_cell(table, MARK_ROW_OFFSET + 1, 1, fixed.get("合同名称", ""))
    fill_value_cell(table, MARK_ROW_OFFSET + 2, 1, sender)

    fill_cell(table, MARK_ROW_OFFSET + 3, 0, f"{buyer}采购申请处室项目责任人")
    fill_value_cell(table, MARK_ROW_OFFSET + 3, 1, contract_manager)
    fill_cell(table, MARK_ROW_OFFSET + 4, 0, f"{buyer}商务合同经理及接货人电话")
    fill_value_cell(table, MARK_ROW_OFFSET + 4, 1, f"接货人电话:{receiver} {receiver_phone}（合同经理:{contract_manager}）")

    fill_value_cell(table, MARK_ROW_OFFSET + 5, 1, ship_from_text)
    fill_value_cell(table, MARK_ROW_OFFSET + 6, 1, destination)
    fill_value_cell(table, MARK_ROW_OFFSET + 7, 1, str(box_no))
    fill_value_cell(table, MARK_ROW_OFFSET + 7, 3, str(total_count))
    fill_value_cell(table, MARK_ROW_OFFSET + 8, 1, box_first_value(items, "包装尺寸"))
    fill_value_cell(table, MARK_ROW_OFFSET + 8, 3, box_first_value(items, "毛重/净重"))
    fill_value_cell(table, MARK_ROW_OFFSET + 9, 1, unique_join([item["合同序号"] for item in items]))
    fill_value_cell(table, MARK_ROW_OFFSET + 9, 3, box_item_name(items))


def render_marks(template_path: Path, output_path: Path, fixed: dict[str, str], box_groups: OrderedDict[int, list[dict[str, str]]], max_box: int) -> dict[str, object]:
    if not box_groups:
        raise ValueError("发货清单中未找到可展开的箱号。")

    template_doc = Document(template_path)
    template_table = deepcopy(template_doc.tables[0]._tbl)
    doc = Document(template_path)
    clear_document_body(doc)
    tighten_page_margins(doc)

    for _box_no in box_groups.keys():
        table_element = deepcopy(template_table)
        append_table_before_section(doc, table_element)

    for table, (box_no, items) in zip(doc.tables, box_groups.items()):
        insert_title_row(table)
        fill_mark_table(table, fixed, box_no, max_box, items)
        apply_table_font(table)

    for table in doc.tables[1:]:
        set_page_break_before_table(table)

    doc.save(output_path)
    return {
        "output": str(output_path),
        "marks": len(box_groups),
        "max_box": max_box,
        "first_box_sequences": unique_join([item["合同序号"] for item in box_groups.get(1, [])]),
        "first_box_total_count": max_box,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="按中核模板生成每箱一页唛头。")
    parser.add_argument("contract_folder", help="合同文件夹路径。")
    parser.add_argument("--template", required=True, help="中核唛头模板 docx 路径。")
    parser.add_argument("--output-folder", help="输出文件夹路径。")
    args = parser.parse_args()

    contract_folder = Path(args.contract_folder).resolve()
    template_path = Path(args.template).resolve()
    info_path = contract_folder / INFO_TABLE_NAME
    if not info_path.exists():
        raise SystemExit(f"未找到信息搜集表：{info_path}")
    if not template_path.exists():
        raise SystemExit(f"未找到唛头模板：{template_path}")

    fixed, batch = read_info_table(info_path)
    batch_name = batch.get("批次") or "第一批"
    shipment_name = batch.get("发货清单文件") or "发货清单（第一批）.xlsx"
    shipment_path = contract_folder / shipment_name
    if not shipment_path.exists():
        raise SystemExit(f"未找到发货清单：{shipment_path}")

    output_folder = (
        Path(args.output_folder).resolve()
        if args.output_folder
        else contract_folder / DEFAULT_OUTPUT_FOLDER.format(batch=batch_name)
    )
    output_folder.mkdir(parents=True, exist_ok=True)
    output_path = output_folder / f"唛头（每箱1张）-{batch_name}-待确认.docx"

    box_groups, max_box = group_items_by_box(read_shipment_rows(shipment_path))
    result = render_marks(template_path, output_path, fixed, box_groups, max_box)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
